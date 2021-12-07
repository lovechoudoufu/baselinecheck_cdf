#!/usr/bin/python3
# -*- coding:utf-8 -*-

from docx import Document
import re
import os

workpath = os.getcwd()  # 同其他


class CheckWin:   # windows的把raw中的命令执行原始结果，一项项进行判断是否合规，然后输出成word报告。
    checkresult = [''] * 29  # 定义了一个存放是否合格的结果列表

    def __init__(self, ipfilepath, theorder, ipfile):
        self.filelist = []   # filelist 以list列表的形式用来存放raw内容的。['第一行','第二行','第三行','第四行']
        self.locatelist = []  # 分割行 ['3','8','15','29']
        self.resultlist = [''] * 29  # 存放检查结果的list列表，内容对应到word表上
        self.checklist = []
        self.IP = '0.0.0.0'
        self.file = ipfilepath  # Raw文件比如：raw\windows\999999999\192.168.0.1
        self.theorder = theorder
        self.ipfile = ipfile    # 命令执行原始结果保存的文件名比如：192.168.0.1

        try:
            f = open(self.file, 'r')    # 以读的方式打开检查脚本文件  # 打开命令结果：raw\windows\999999999\192.168.0.1
        except:
            f = open(self.file, 'r',encoding='utf-8')  # 如果因为文件编码问题导致错误就用utf-8打开
        count = 0  # 记录行数的
        for line in f:
            count += 1
            self.filelist.append(line)   # filelist 以list列表的形式用来存放raw内容的
            if 'IsCheckLine' in line:  # 分割行 确定检查结果的行数，结果中带有IsCheckLine的都是一个新的开始，两个IsCheckLine之间的是上一个的结果
                self.locatelist.append(count)  #  locatelist 这个list列表用来存放IsCheckLine分割行的
        self.locatelist.append(count)
        f.close()

    def read_systeminfo(self):  # 获取系统基本信息   从systeminfo命令的执行结果中获取
        for i in range(self.locatelist[0], self.locatelist[1] - 1):   # 上一个分割行到下一个分割行-1 是一堆检查内容
            # print(self.filelist[i])
            if 'OS 名称' in str(self.filelist[i]):  # if判断这行
                self.resultlist[1] = self.filelist[i].split(':')[1].strip(' ')  #  获取os 名称这行后面的东西，赋值给resultlist[1]，后面填写到word报告中

    def read_ipconfig(self):   # 获取ip地址信息   从ipconfig的命令执行结果中获取
        iplist = []
        for i in range(self.locatelist[1], self.locatelist[2] - 1):   # 上一个分割行到下一个分割行-1 是一堆检查内容
            if 'IPv4' in str(self.filelist[i]):    #  从结果中找到IPv4这行
                iplist.append(re.findall(r"(\d*\.\d*\.\d*\.\d*)", self.filelist[i])[0])  # re库按照ip的正则语法匹配出ip地址
                self.resultlist[0] = iplist   # 将ip地址保存到resultlist[0]里面，放到后面word报告中
        self.IP = self.resultlist[0]

    def read_users(self):   # 获取系统中的用户名  从net user和net user administrtor的命令执行结果中获取
        for i in range(self.locatelist[3], self.locatelist[4] - 1):
            if '用户名' in str(self.filelist[i]):   #找到用户名所在位置
                name = self.filelist[i].split('               ')[1].strip(' ').strip('\n') + ','
                #去掉前面空格（有多少代码中对应多少）后面第一个位置就是用户名，再去掉后面空格（防止后面有空格）和换行拼接上’，‘
                if 'Yes' in str(self.filelist[i + 5]): #是否开启在用户名之后的第5行，如果是yes就是开启
                    self.resultlist[4] += name       #然后把name赋值给它，resultlist[4]里面保存着所有开启的用户名
        if 'Administrator' in self.resultlist[4]:   #  判断administrator用户是否禁用或者已经重命名掉，启用的话不合格，正常基线要求是不启用administrator账户
            self.resultlist[2] = '(x)未更改'  #  不合格的标记一个(x)
        else:
            self.resultlist[2] = '已禁用或已更改'
        if 'Guest' in self.resultlist[4]:    # 判断guest用户是否是禁用状态，启用的话不合格
            self.resultlist[3] = '(x)未禁用'  # 不合格标记一个(x)
        else:
            self.resultlist[3] = '已禁用'

    def read_secedit(self):   # 查看系统策略是否合格，包括密码策略、等
        audit = {'0': '(x)无审核', '1': '(x)成功', '2': '(x)失败', '3': '成功、失败', } # 只有成功和失败都审核时候才合格
        for i in range(self.locatelist[2], self.locatelist[3] - 1):  # 一行一行的拿出两个分割行之间的命令执行结果
            if '=' in self.filelist[i]:  # 判断这一行里面有没有=号
                str = self.filelist[i].split('=')[0].strip(' ') #  以 = 为分割，把一行分成key和值，把key给str参数
                res = self.filelist[i].split('=')[1].strip(' ').strip('\n') # 以 = 为分割，把一行分成key和值，把值给res参数
                # 例如第一个，判断密码复杂度，判断一下这一行的key是不是PasswordComplexity，是的话进入if中判断下值是否符合基线的要求，符合就合格，不符合就标记(x)为不合格
                if str == 'PasswordComplexity':  # 密码必须符合复杂性要求  密码相关策略：https://docs.microsoft.com/en-us/openspecs/windows_protocols/ms-gpsb/0b40db09-d95d-40a6-8467-32aedec8140c
                    # 为0时候没有开启密码复杂度限制，为不合格，非0时候为合格
                    if res == '0': # 该项要求为0时候不符合
                        self.resultlist[5] = '(x)未启用密码复杂度要求'
                    else:
                        self.resultlist[5] = '已启用'
                elif str == 'MinimumPasswordLength':  # 密码长度最小值
                    # 为0时候未设置密码长度限制，不合格；0<密码长度<8时候也不合格；密码长度大于8才合格
                    if int(res) == 0:
                        self.resultlist[6] = "(x)不限制密码长度，值：" + res
                    elif 0 < int(res) < 8:
                        self.resultlist[6] = "(x)长度小于8位，值："+res
                    else:
                        self.resultlist[6] = res + " 符合"
                elif str == 'MaximumPasswordAge':  # 密码最长使用期限
                    # 大于90天不合格，为-1是不设置过期时间也不合格
                    if int(res) > 90:
                        self.resultlist[7] = "(x)时间大于90天，值："+res
                    elif int(res) == -1:
                        self.resultlist[7] = "(x)密码永不过期，值："+res
                    else:
                        self.resultlist[7] = res+" 符合"
                elif str == 'PasswordHistorySize':  # 强制历史密码的相似，不与以前使用的密码重合
                    # 检查与以前密码的相似性
                    if int(res) == 0:
                        self.resultlist[8] += "(x)不检查历史密码记录，值："+res
                    elif 0 < int(res) < 5:
                        self.resultlist[8] += "(x)数值小于5，值："+res
                    else:
                        self.resultlist[8] = res + " 符合"
                elif str == 'LockoutBadCount':  # 密码错误多少次后账户锁定，0为不锁定不合规
                    # 登录密码错误多少次后锁定登录账号，0为未设置不合规，大于6次不合规
                    if int(res) == 0:
                        self.resultlist[9] = "(x)未开启密码错误多次锁定账户的策略"
                    elif int(res) > 6:
                        self.resultlist[9] = "(x)错误次数大于6次"
                    else:
                        self.resultlist[9] = res + "符合"

                elif str == 'SeRemoteShutdownPrivilege':  # 允许远程强制关机服务器，例如3389远程桌面上关闭服务器，应只允许administrators组的用户关闭，其他不可
                    self.resultlist[10] = res.replace('*S-1-5-32-544', 'Administrators').replace('*S-1-1-0',
                                                                                                 'Everyone').replace(
                        '*S-1-5-32-545', 'Users').replace('*S-1-5-32-547', 'Users').replace('S-1-5-32-551',
                                                                                            'Backup Operators')
                    if res != '*S-1-5-32-544':
                        self.resultlist[10] += "(x)允许非Administrators组用户远程关闭计算机"  # 只有管理员组才能关，其他组都不能关
                elif str == 'SeTakeOwnershipPrivilege':
                    self.resultlist[11] = res.replace('*S-1-5-32-544', 'Administrators').replace('*S-1-1-0',
                                                                                                 'Everyone').replace(
                        '*S-1-5-32-545', 'Users').replace('*S-1-5-32-547', 'Users').replace('S-1-5-32-551',
                                                                                            'Backup Operators')
                    if res != '*S-1-5-32-544':
                        self.resultlist[11] += "(x)允许非Administrators组用户获取文件对象所有权"  # 同上
                # 以下对应着audit的字典关系：audit = {'0': '(x)无审核', '1': '(x)只成功', '2': '(x)只失败', '3': '成功和失败', }，只有为3，同时审核成功和失败时候才算是合格
                elif str == 'AuditPolicyChange':
                    self.resultlist[12] = audit[res]
                elif str == 'AuditAccountLogon':
                    self.resultlist[13] = audit[res]
                elif str == 'AuditObjectAccess':
                    self.resultlist[14] = audit[res]
                elif str == 'AuditProcessTracking':
                    self.resultlist[15] = audit[res]
                elif str == 'AuditDSAccess':
                    self.resultlist[16] = audit[res]
                elif str == 'AuditPrivilegeUse':
                    self.resultlist[17] = audit[res]
                elif str == 'AuditSystemEvents':
                    self.resultlist[18] = audit[res]
                elif str == 'AuditAccountManage':
                    self.resultlist[19] = audit[res]

    def read_oth(self):  # 检查其他内容
        for i in range(self.locatelist[4], self.locatelist[5] - 1):
            if '自动播放' in str(self.filelist[i]):  # 是否关闭自动播放策略（例如插入u盘后，盘中文件自动播放）
                if 'oxff' in str(self.filelist[i + 1]):
                    self.resultlist[22] = '已关闭'
                else:
                    self.resultlist[22] = '(x)开启了自动播放功能'
            if '防火墙状态' in str(self.filelist[i]): # 是否开启系统防火墙
                if '0x0' in str(self.filelist[i + 1]):
                    self.resultlist[21] = '(x)关闭了防火墙功能'
                else:
                    self.resultlist[21] = '已开启'
            if '远程桌面' in str(self.filelist[i]): # 是否关闭windows远程桌面功能
                if '0x0' in str(self.filelist[i + 1]):
                    self.resultlist[25] = '(x)开启了远程桌面功能'
                else:
                    self.resultlist[25] = '已关闭'
            if '3389端口' in str(self.filelist[i]): # 远程桌面端口是否为默认的3389端口
                if '0xd3d' in str(self.filelist[i + 1]):
                    self.resultlist[26] = '(x)远程桌面功能为默认端口'
                else:
                    self.resultlist[26] = '非默认端口'
            if '分区共享' in str(self.filelist[i]): # 判断是否开启文件共享
                if '0x0' in str(self.filelist[i + 1]):
                    self.resultlist[24] += "分区共享：已关闭\n"
                else:
                    self.resultlist[24] += "(x)分区共享：未关闭\n"
            if 'ADMIN共享' in str(self.filelist[i]):# 判断是否开启文件共享
                if '0x0' in str(self.filelist[i + 1]):
                    self.resultlist[24] += "ADMIN共享：已关闭\n"
                else:
                    self.resultlist[24] += "(x)ADMIN共享：未关闭\n"
            if 'IPC共享' in str(self.filelist[i]):# 判断是否开启文件共享
                if '0x1' in str(self.filelist[i + 1]):
                    self.resultlist[24] += "IPC共享：已关闭\n"
                else:
                    self.resultlist[24] += "(x)IPC共享：未关闭\n"
            # 从这往下是日志
            if '应用日志文件大小' in str(self.filelist[i]):# 判断应用日志的大小设置是否合格
                self.resultlist[20] += '应用日志文件大小：' + self.filelist[i + 1].split('REG_DWORD')[1].strip(' ').strip('\n')
                if int(self.filelist[i + 1].split('REG_DWORD')[1].strip(' '), 16) < 41943040:
                    self.resultlist[20] += '(x)\n'
                else:
                    self.resultlist[20] += '\n'
                if '0x0' in self.filelist[i + 3]:
                    self.resultlist[20] += '到事件日志最大大小时：' + '按需要覆盖事件（旧事件优先）\n'
                else:
                    self.resultlist[20] += '到事件日志最大大小时：' + '(x)\n'
            if '安全日志文件大小' in str(self.filelist[i]): # 判断安全日志的大小设置是否合格
                self.resultlist[20] += '安全日志文件大小：' + self.filelist[i + 1].split('REG_DWORD')[1].strip(' ').strip('\n')
                if int(self.filelist[i + 1].split('REG_DWORD')[1].strip(' '), 16) < 41943040:
                    self.resultlist[20] += '(x)\n'
                else:
                    self.resultlist[20] += '\n'
                if '0x0' in self.filelist[i + 3]:  # 日志达到最大值时候怎么着，0是覆盖旧日志，符合
                    self.resultlist[20] += '到事件日志最大大小时：' + '按需要覆盖事件（旧事件优先）\n'
                else:
                    self.resultlist[20] += '到事件日志最大大小时：' + '(x)\n'
            if '系统日志文件大小' in str(self.filelist[i]): # 判断系统日志的大小设置是否合格
                self.resultlist[20] += '系统日志文件大小：' + self.filelist[i + 1].split('REG_DWORD')[1].strip(' ').strip('\n')
                if int(self.filelist[i + 1].split('REG_DWORD')[1].strip(' '), 16) < 41943040:
                    self.resultlist[20] += '(x)\n'
                else:
                    self.resultlist[20] += '\n'
                if '0x0' in self.filelist[i + 3]:
                    self.resultlist[20] += '到事件日志最大大小时：' + '按需要覆盖事件（旧事件优先）'
                else:
                    self.resultlist[20] += '到事件日志最大大小时：' + '(x)'
            # 获取windows的最新更新系统补丁的时间
            if 'WindowsUpdate' in str(self.filelist[i]):
                self.resultlist[23] = str(self.filelist[i]).split(':')[1]

    def ipdocx(self):      # 输出单个ip的检查报告
        doc = Document(workpath + '\\models\\Wordtemplate\\Windows.docx')     # 打开word模板
        doc.paragraphs[0].text = str(self.IP) + ':'     # word的第一行，用来写ip
        t = doc.tables[0]  # 打开tables（表格）
        count = 0
        for row in t.rows:  # 一行一行的操作这个表格
            if count == 0:
                count += 1
                pass
            elif count == 29:
                break
            else:
                # row.cells[2].text 是第二列，resultlist列表里面的内容写到对性的行上
                row.cells[2].text = str(self.resultlist[count - 1]) # 把上面的resultlist列表中的结果按照次序写入到word中
                count += 1
        try:
            docxpath = workpath + '\\models\\Result\\' + self.theorder  # word要保存的目录
            if not os.path.exists(docxpath):
                os.mkdir(docxpath)
            doc.save(docxpath + '\\' + self.ipfile + '.docx')    # 保存写好的word到新的文件
        except IOError:
            print("please close word")

    def linenook(self):     # 排查一下不符合的项目
        for i in range(0, len(self.resultlist)):
            if '(x)' in str(self.resultlist[i]):  # 列表中一个一个的看里面有没有(x)，带(x)的不符合
                self.checklist.append(i)  # checklist 不符合的放到这个列表里

    def result(self):
        for i in self.checklist:
            self.checkresult[int(i)] += str(self.IP) + '\n'  # 不符合的把ip地址保存上

    def resultdocx(self):  # 输出整体的不符合报告，哪个项目有哪些ip不符合
        doc = Document(workpath + '\\models\\Wordtemplate\\Windows-result.docx')
        t = doc.tables[0]
        count = 0
        for row in t.rows:
            if count == 0:
                count += 1
                pass
            elif count == 29:
                break
            else:
                row.cells[2].text = str(self.checkresult[count + 1][:-1])
                count += 1
        try:
            docxpath = workpath + '\\models\\Result\\' + self.theorder
            if not os.path.exists(docxpath):
                os.mkdir(docxpath)
            doc.save(docxpath + '\\all-result.docx')
        except IOError:
            print("please close word")

    def check(self): # check函数里面是调用的上面的判断是否合规的函数。
        try:
            self.read_systeminfo()
        except Exception as e:
            pass
        try:
            self.read_ipconfig()
        except Exception as e:
            pass
        try:
            self.read_secedit()
        except Exception as e:
            pass
        try:
            self.read_users()
        except Exception as e:
            pass
        try:
            self.read_oth()
        except Exception as e:
            pass
        try:
            self.ipdocx()
        except Exception as e:
            pass
        try:
            self.linenook()
        except Exception as e:
            pass
        try:
            self.result()
        except Exception as e:
            pass
        try:
            self.resultdocx()
        except Exception as e:
            pass
