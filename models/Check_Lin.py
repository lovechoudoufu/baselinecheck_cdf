#!/usr/bin/python3
# -*- coding:utf-8 -*-

from docx import Document
import re
import os
#导入所需要的库

workpath = os.getcwd() #获取当前运行文档的路径


class CheckLin:     # 定义CheckLin这个类
    checkresult = [''] * 32     # 一维数组，里面有32个元素，赋值给checkresult

    def __init__(self, ipfilepath, theorder, ipfile):     # 初始化这些函数
        self.filelist = []
        self.locatelist = []
        self.resultlist = [''] * 22
        self.checklist = []
        self.IP = ipfile.split('/')[-1]    # 用/对获取到的IP进行分片，-1是默认的，就是分割所有
        self.file = ipfilepath      # 获取存放Linux检查脚本文件夹路径
        self.theorder = theorder     # 从数据库中获取theorder编号
        self.ipfile = ipfile     # 获取IP文件

        try:
            f = open(self.file, 'r',encoding='utf-8')    # 以可读的方式打开检查脚本文件，utf-8编码格式打开
        except:
            f = open(self.file, 'r')  # 如果因为文件编码问题导致错误就不加utf-8
        count = 0      # 记录行数
        for line in f:
            count += 1
            self.filelist.append(line)   # filelist 以list列表的形式用来存放raw内容的
            if 'IsCheckLine' in line:    # 分割行 确定检查结果的行数，结果中带有IsCheckLine的都是一个新的开始，两个IsCheckLine之间的是上一个的结果
                self.locatelist.append(count)    #  locatelist 这个list列表用来存放IsCheckLine分割行的
        self.locatelist.append(count)
        f.close()

    def inetip(self):  # 获取系统基本信息
        iplist = []
        for i in range(self.locatelist[0], self.locatelist[1] - 1):    # 上一个分割行到下一个分割行-1 是一堆检查内容
            if 'inet' in str(self.filelist[i]).lower():     # lower是把这行内容转换成小写
                ip = re.match(r'.*inet (.*)/.*', str(self.filelist[i]))[1]    # 正则语法匹配出inet 后/前的ip地址
                iplist.append(ip)  # 存放到iplist这个列表里
            else:
                pass
        self.resultlist[1] = iplist  # 放到resultlist中第一个位置

    def linuxcore(self): #定义linuxcore这个函数
        for i in range(self.locatelist[1], self.locatelist[2] - 1):     # 分割行 确定检查结果的行数，结果中带有IsCheckLine的都是一个新的开始，两个IsCheckLine之间的是上一个的结果
            if str(self.filelist[i]):
                self.resultlist[2] = str(self.filelist[i]).strip('\n')    # 把Linux系统内核信息存放到resultlist第二个位置中

    def linuxuser(self): # 获取uid为0的用户名和可登录的用户名
        user0uid = []
        usercanlogin = []
        for i in range(self.locatelist[2], self.locatelist[3] - 1):
            if '0' == str(self.filelist[i]).split(' ')[1]:  # 看一下第一位是不是0，uid为0的话就是root管理员权限
                user0uid.append(str(self.filelist[i]).split(' ')[0])  # 把uid为0的用户名加到user0uid列表里
            if 'bash' in str(self.filelist[i]).lower().split(' ')[3]:    # 看一下第三位是不是带bash的，找到bash所在位置，带bash的是可登录账户
                usercanlogin.append(str(self.filelist[i]).split(' ')[0])  # 把可登录的额用户名加到usercanlogin这个列表中
        if len(user0uid):  # 如果user0uid列表不为空，就把列表给resultlist的3位置
            self.resultlist[3] = user0uid
        else:
            self.resultlist[3] = 'ok'
        self.resultlist[4] = usercanlogin

    def linuxgroup(self): # 检查root用户组的用户和可通过sudo su提权到root权限的用户
        self.resultlist[5] = '无'
        templist = []
        for i in range(self.locatelist[3], self.locatelist[4] - 1):
            if 'rootgroup' in str(self.filelist[i]):  # 用户组在root里面的用户：cat /etc/group
                self.resultlist[5] = str(self.filelist[i]).split(':')[1].strip('\n')
            if 'ALL=' in str(self.filelist[i]).upper() and 'ROOT' not in str(self.filelist[i]).upper():  # 那些用户可以通过sudo su命令提升权限到root权限，：cat /etc/sudoers
                templist.append(str(self.filelist[i]).split(' ')[0])
        if len(templist) == 0:
            self.resultlist[6] = "无"

    def passwdcheck(self):  # 密码策略检查
        denynum = 4
        unlocktimenum = 601
        self.resultlist[11] = '(x)'
        for i in range(self.locatelist[4], self.locatelist[5] - 1):
            if 'PASS_MAX_DAYS' in str(self.filelist[i]).upper():  # 检查密码口令强制更改周期的天数
                num1 = (re.split('\t| ', str(self.filelist[i]))[1]).strip('\n')
                if int(num1) > 90: # 天数大于90天不合格
                    self.resultlist[7] = '(x) ' + num1
                else:
                    self.resultlist[7] = num1
            if 'PASS_MIN_DAYS' in str(self.filelist[i]).upper(): # 口令更改最小时间间隔的天数
                num2 = (re.split('\t| ', str(self.filelist[i]))[1]).strip('\n')
                if int(num2) != 0:  # 0是没有限制，所以不符合
                    self.resultlist[8] = '(x)' + num2
                else:
                    self.resultlist[8] = num2
            if 'PASS_WARN_AGE' in str(self.filelist[i]).upper():  # 口令过期警告时间天数
                num3 = (re.split('\t| ', str(self.filelist[i]))[1]).strip('\n') # 以制表符或者空格为分隔，然后一行分成数组，取数组的1位
                if int(num3) < 3: # 天数小于3不合格
                    self.resultlist[9] = '(x)' + num3
                else:
                    self.resultlist[9] = num3
            if 'PASS_MIN_LEN' in str(self.filelist[i]).upper():  # 口令最小长度
                num4 = (re.split('\t| ', str(self.filelist[i]))[1]).strip('\n')
                if int(num4) < 8:  # 小于8不合格
                    self.resultlist[10] = '(x) ' + num4
                else:
                    self.resultlist[10] = str(self.filelist[i]).split('  ')[1]
            if 'unlock_time' in str(self.filelist[i]).lower():    # 如果找到锁定时间这个命令执行结果，就转化成小写形式
                try:
                    denynum = re.match(r'deny=(\d+)', str(self.filelist[i]).lower())[1]
                    unlocktimenum = re.match(r'unlock_time=(\d+)', str(self.filelist[i]).lower())[1]
                except:
                    pass
                if int(denynum) > 5 or int(unlocktimenum) < 600:
                    self.resultlist[11] = '(x)'
                else:
                    self.resultlist[11] = 'ok'

    def umask(self): # 查看umask
        for i in range(self.locatelist[5], self.locatelist[6] - 1):
            if 'umask' in str(self.filelist[i]):
                self.resultlist[12] = 'ok'

    def checklog(self):  # 检查syslog日志服务器是否开启、btmp、lastlog、secure、wtmp日志都是否存在
        self.resultlist[13] = '(x)'
        self.resultlist[14] = '(x)'
        btmp_status = 0
        lastlog_status = 0
        secure_status = 0
        wtmp_status = 0
        for i in range(self.locatelist[6], self.locatelist[7] - 1):
            if '/usr/sbin/rsyslogd' in str(self.filelist[i]) or '/usr/sbin/syslogd' in str(self.filelist[i]):
                self.resultlist[13] = 'running'
            if '/sbin/auditd' in str(self.filelist[i]):
                self.resultlist[14] = 'running'
            if 'btmp' in str(self.filelist[i]):
                btmp_status = 1
            if 'lastlog' in str(self.filelist[i]):
                lastlog_status = 1
            if 'secure' in str(self.filelist[i]):
                secure_status = 1
            if 'wtmp' in str(self.filelist[i]):
                wtmp_status = 1
        if btmp_status * lastlog_status * secure_status * wtmp_status == 0:
            self.resultlist[15] = '(x)'
        else:
            self.resultlist[15] = 'ok'

    def rootlogin(self):
        self.resultlist[16] = '(x)'
        for i in range(self.locatelist[7], self.locatelist[8] - 1):
            # permitrootlogin 在/etc/ssh/sshd_config是否设置，设置了yes才符合，不设置不符合
            if 'permitrootlogin' in str(self.filelist[i]).lower() and 'yes' in str(self.filelist[i]).lower():
                self.resultlist[16] = 'ok'

    def checktelnet(self): # 检查是否开启了telnet服务，如果开启了就是不合格
        self.resultlist[17] = 'ok'
        for i in range(self.locatelist[8], self.locatelist[9] - 1):
            if 'telnet' in str(self.filelist[i]).lower() and 'grep telnet' not in str(self.filelist[i]).lower():
                self.resultlist[17] = '(x)'

    def checkrsalogin(self): # 检查是否开启了ssh的rsa公钥登录，即免密登录，开启了的话就不符合
        templist = []
        for i in range(self.locatelist[9], self.locatelist[10] - 1):
            if 'authorized_keys' in str(self.filelist[i]).lower():
                templist.append(str(self.filelist[i]).strip('\n'))    #把authorized_keys放到templist列表中并且去掉换行
        if len(templist) == 0:
            self.resultlist[18] = 'ok'
        else:
            self.resultlist[18] = '(x)' + str(templist)

    def checkweb(self):  # 检查web容器比如tomact这些的运行权限，如果是root权限则不合格
        self.resultlist[19] = 'ok'
        for i in range(self.locatelist[10], self.locatelist[11] - 1):
            if 'grep' not in str(self.filelist[i]).lower():
                if re.search(r'tomcat|weblogic|httpd|jboss|nginx', str(self.filelist[i]), re.I):    # re.I：忽略大小写
                    if str(self.filelist[i]).lower().split(' ')[0] == 'root':    #判断权限是否为root
                        self.resultlist[19] = '(x)'

    def checkdist(self):    # 检查系统硬盘的使用
        self.resultlist[20] = 'ok'
        for i in range(self.locatelist[11], self.locatelist[12] - 1):     # 把这段内容一行一行读取进行for循环判断
            try:
                disttmp = re.findall('(\d+)%', str(self.filelist[i]))[0]    # 找到0位置上的内存使用率
                if int(disttmp) > 80:    # 判断使用率是否>80%
                    self.resultlist[20] = '(x)'    # 大于不合格，不大于是合格
            except:
                pass  # one line Filesystem               Size  Used Avail Use% Mounted on

    def checkfile(self):    # 检查几个重要文件的权限
        self.resultlist[21] = 'ok'
        templist = []    # 定义存放结果的列表
        for i in range(self.locatelist[12], self.locatelist[13] - 1):    # 把确定范围里面的内容一行一行读取出来
            # r 读取  w 写入 x 执行 - --- --- ---  第一个-是文件类型，第二个---是所有者的，第二个---是所有者的组的，第三个---是其他组用户的
            if 'passwd' in str(self.filelist[i]) and '-rw-r--r--' not in str(self.filelist[i]):    # 判断一下passwd权限是否是可修改的
                templist.append(str(self.filelist[i]))
            if 'shadow' in str(self.filelist[i]) and '----------' not in str(self.filelist[i]):
                templist.append(str(self.filelist[i]))
            if 'sudoers' in str(self.filelist[i]) and '-r--r-----' not in str(self.filelist[i]):
                templist.append(str(self.filelist[i]))
            if 'group' in str(self.filelist[i]) and '-rw-r--r--' not in str(self.filelist[i]):
                templist.append(str(self.filelist[i]))
            if 'profile' in str(self.filelist[i]) and '-rw-r--r--' not in str(self.filelist[i]):
                templist.append(str(self.filelist[i]))
        if len(templist) != 0:
            self.resultlist[21] = '(x)' + str(templist)

    # def testresult(self):
    # print(self.resultlist)

    def ipdocx(self):    # 输出单个ip的检查报告
        # print(self.resultlist)
        doc = Document(workpath + '\\models\\Wordtemplate\\linux.docx')    # 打开Linux模板
        doc.paragraphs[0].text = str(self.IP) + ':'    #  Linux的第一行，用来写ip
        t = doc.tables[0]    # 打开tables（表格）
        count = 0
        for row in t.rows:    # 一行一行的操作这个表格
            if count == 0:
                count += 1
                pass
            elif count == 22:
                break
            else:
                # row.cells[2].text 是第二列，resultlist列表里面的内容写到对性的行上
                row.cells[2].text = str(self.resultlist[count])
                count += 1
        try:
            docxpath = workpath + '\\models\\Result\\' + self.theorder    # Linux报告要保存的目录
            if not os.path.exists(docxpath):
                os.mkdir(docxpath)
            doc.save(docxpath + '\\' + self.ipfile + '.docx')    # 保存写好的Linux到新的文件
        except IOError:
            print("please close word")

    def lineno(self):    # 排查一下不符合的项目
        for i in range(0, 22):
            if '(x)' in str(self.resultlist[i]):    # 列表中一个一个的看里面有没有(x)，带(x)的不符合
                self.checklist.append(i)     # checklist 不符合的放到这个列表里

    def result(self):
        for i in self.checklist:
            self.checkresult[int(i)] += str(self.IP) + '\n' # 不符合的把ip地址保存上

    def resultdocx(self):
        doc = Document(workpath + '\\models\\Wordtemplate\\Linux-result.docx') # 用Document函数打开报告模板
        t = doc.tables[0] # 打开报告的表格
        count = 0
        for row in t.rows:
            if count == 0: # 第0行跳过
                count += 1
                pass
            elif count == 16: # 第16行结束，因为Linux-result.docx总共16行
                break
            else:
                row.cells[2].text = str(self.checkresult[count + 6][:-1]) #总表里面从第七个开始
                count += 1
        try:
            docxpath = workpath + '\\models\\Result\\' + self.theorder # 保存路径
            if not os.path.exists(docxpath):
                os.mkdir(docxpath)
            doc.save(docxpath + '\\all-result.docx') # 保存成文件
        except IOError:
            print("please close word")

    def check(self):
        try:
            self.inetip()
        except Exception as e:
            pass
        try:
            self.linuxcore()
        except Exception as e:
            pass
        try:
            self.linuxuser()
        except Exception as e:
            pass
        try:
            self.linuxgroup()
        except Exception as e:
            pass
        try:
            self.passwdcheck()
        except Exception as e:
            pass
        try:
            self.umask()
        except Exception as e:
            pass
        try:
            self.checklog()
        except Exception as e:
            pass
        try:
            self.rootlogin()
        except Exception as e:
            pass
        try:
            self.checktelnet()
        except Exception as e:
            pass
        try:
            self.checkrsalogin()
        except Exception as e:
            pass
        try:
            self.checkweb()
        except Exception as e:
            pass
        try:
            self.checkdist()
        except Exception as e:
            pass
        try:
            self.checkfile()
        except Exception as e:
            pass
        try:
            self.ipdocx()
        except Exception as e:
            pass
        try:
            self.lineno()
        except Exception as e:
            pass
        try:
            self.result()
        except Exception as e:
            pass
        try:
            self.resultdocx()
        except Exception as e:
            pass
